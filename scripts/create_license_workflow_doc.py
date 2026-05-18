from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/generated/SMART_SnippetFlow_Lizenz_Workflow_Dokumentation.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 37, 55)
MUTED = RGBColor(93, 108, 126)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)
        borders.append(tag)


def set_table_width(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text: str, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_step(doc: Document, number: int, title: str, detail: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    add_run(paragraph, f"{title}: ", bold=True)
    paragraph.add_run(detail)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    set_table_width(table, [1.85, 4.65])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Body"
            set_cell_margins(cell)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], LIGHT_BLUE)
        for paragraph in header_cells[index].paragraphs:
            paragraph.style = "Table Header"
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.style = "Table Body"
            set_cell_margins(cells[index])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    table_header = styles.add_style("Table Header", 1)
    table_header.font.name = "Calibri"
    table_header.font.size = Pt(10)
    table_header.font.bold = True
    table_header.font.color.rgb = INK
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.1

    table_body = styles.add_style("Table Body", 1)
    table_body.font.name = "Calibri"
    table_body.font.size = Pt(10)
    table_body.font.color.rgb = INK
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.1


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = doc.styles["Normal"]
    run = footer.add_run("SMART SnippetFlow Lizenzsystem")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_footer(doc)

    doc.add_paragraph("SMART SnippetFlow Lizenzsystem", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Workflow-Dokumentation: Landingpage, Stripe Checkout, Supabase-Lizenzdatenbank und Desktop-Aktivierung")

    doc.add_heading("1. Zweck und Systemüberblick", level=1)
    doc.add_paragraph(
        "Diese Dokumentation beschreibt den End-to-End-Prozess für Kauf, Lizenzanlage, App-Download und Aktivierung "
        "der Einzelplatzlizenz von SMART SnippetFlow. Stripe ist die Zahlungsquelle, Supabase hält Lizenzstatus und "
        "Geräteaktivierungen, und die Desktop-App prüft den Lizenzstatus über eingeschränkte Supabase-RPCs."
    )
    add_key_value_table(
        doc,
        [
            ("Zahlungsanbieter", "Stripe Checkout, Subscription, Rechnung und Webhook-Ereignisse"),
            ("Lizenzdatenbank", "Supabase Postgres mit Tabellen fuer Kunden, Lizenzen, Aktivierungen und Audit-Logs"),
            ("Desktop-App", "Electron-App mit lokaler Lizenzcache-Datei und Supabase-Aktivierung"),
            ("Aktueller Modus", "Einzelplatzlizenz, 1 Sitz, 1 aktives Gerät"),
            ("Erweiterbar für", "Cloud-Sync, Teams, mehrere Geräte und Kundenportal"),
        ],
    )

    doc.add_heading("2. Beteiligte Komponenten", level=1)
    add_matrix_table(
        doc,
        ["Komponente", "Aufgabe", "Wichtige Daten"],
        [
            ["Landingpage", "Startet den Kaufprozess per POST-Aufruf an die Checkout Function.", "VITE_CHECKOUT_FUNCTION_URL"],
            ["Supabase Function create-checkout-session", "Erstellt eine frische Stripe Checkout Session.", "STRIPE_SECRET_KEY, STRIPE_PRICE_ID"],
            ["Stripe Checkout", "Verarbeitet Zahlung und Subscription.", "Checkout Session, Customer, Subscription"],
            ["Supabase Function stripe-webhook", "Verifiziert Stripe Events und erzeugt/aktualisiert Lizenzen.", "STRIPE_WEBHOOK_SECRET, SERVICE_ROLE_KEY"],
            ["Supabase Datenbank", "Speichert Lizenz, Aktivierungen, Stripe Events und Audit-Verlauf.", "licenses, license_activations"],
            ["Desktop-App", "Aktiviert und prüft die Lizenz über Supabase RPCs.", "Publishable Key, Lizenzschluessel"],
        ],
        [1.55, 3.1, 1.85],
    )

    doc.add_heading("3. End-to-End Workflow", level=1)
    steps = [
        ("Landingpage öffnen", "Der Nutzer besucht die SMART-SnippetFlow-Landingpage und klickt auf Jetzt kaufen oder Lizenz sichern."),
        ("Checkout Function aufrufen", "Der Button ruft per POST die Supabase Function create-checkout-session auf."),
        ("Stripe Checkout erzeugen", "Die Function erstellt mit dem Stripe Secret Key eine neue Checkout Session für den Jahrespreis."),
        ("Zahlung durchführen", "Der Nutzer bezahlt über Stripe Checkout. Stripe verarbeitet Zahlung, Rechnung und Subscription."),
        ("Webhook empfangen", "Stripe sendet checkout.session.completed und spätere Subscription-Events an die Supabase Webhook Function."),
        ("Lizenz erzeugen", "Die Webhook Function prüft die Signatur und schreibt Kunde, Lizenz, Event und Audit-Eintrag in Supabase."),
        ("App herunterladen", "Der Nutzer lädt die macOS- oder Windows-App von der Landingpage oder einem Release-Host herunter."),
        ("App installieren", "Der Nutzer installiert SMART SnippetFlow lokal und startet die Desktop-App."),
        ("Lizenz aktivieren", "Der Nutzer trägt den Lizenzschlüssel in Einstellungen > Lizenz ein und klickt Aktivieren."),
        ("Gerät registrieren", "Die App erzeugt eine lokale Installations-ID, hasht sie und ruft activate_license auf."),
        ("Status speichern", "Die App speichert license_key, activation_id, remote_status und checked_at lokal."),
        ("Status prüfen", "Spätere Prüfungen laufen über refresh_license_activation; Deaktivierung über deactivate_license_activation."),
    ]
    for index, (title, detail) in enumerate(steps, 1):
        add_step(doc, index, title, detail)

    doc.add_heading("4. Supabase-Datenmodell", level=1)
    add_matrix_table(
        doc,
        ["Tabelle", "Inhalt", "Zweck"],
        [
            ["license_customers", "E-Mail, Name, Stripe Customer ID", "Zuordnung von Stripe-Kunden zu Lizenzen"],
            ["licenses", "Lizenzschlüssel, Status, Stripe IDs, Limits", "Zentrale Lizenzprojektion für App und Support"],
            ["license_activations", "Gerätehash, Plattform, App-Version, Aktivierungszeit", "Gerätebindung und Device-Limit-Prüfung"],
            ["stripe_events", "Stripe Event ID, Typ, Payload, Verarbeitungsstatus", "Idempotenz und Fehleranalyse"],
            ["license_audit_log", "Lizenzbezogene Ereignisse und Metadaten", "Support- und Debug-Historie"],
        ],
        [1.75, 2.45, 2.3],
    )

    doc.add_heading("5. Wichtige Supabase RPCs", level=1)
    add_matrix_table(
        doc,
        ["RPC", "Aufrufer", "Funktion"],
        [
            ["activate_license", "Desktop-App", "Prüft Lizenzschlüssel, Status und Gerätelimit; erzeugt eine Aktivierung."],
            ["refresh_license_activation", "Desktop-App", "Aktualisiert last_seen_at und gibt aktuellen Lizenzstatus zurück."],
            ["deactivate_license_activation", "Desktop-App", "Setzt deactivated_at und gibt den Geräteplatz wieder frei."],
        ],
        [2.1, 1.45, 2.95],
    )

    doc.add_heading("6. Relevante Konfiguration", level=1)
    add_matrix_table(
        doc,
        ["Ort", "Variable", "Beschreibung"],
        [
            ["Supabase Function Secrets", "STRIPE_SECRET_KEY", "Stripe Secret Key aus demselben Stripe-Kontext wie der Preis."],
            ["Supabase Function Secrets", "STRIPE_PRICE_ID", "Jahrespreis für SMART SnippetFlow, beginnt mit price_."],
            ["Supabase Function Secrets", "STRIPE_WEBHOOK_SECRET", "Signaturgeheimnis des Stripe Webhook-Ziels, beginnt mit whsec_."],
            ["Supabase Function Secrets", "SUPABASE_SERVICE_ROLE_KEY", "Serverseitiger Supabase-Key für Webhook-Schreibzugriffe."],
            ["Desktop-App .env", "SMART_SNIPPETFLOW_SUPABASE_URL", "Projekt-URL: https://nrnerbyljrvlrximotsj.supabase.co"],
            ["Desktop-App .env", "SMART_SNIPPETFLOW_SUPABASE_ANON_KEY", "Publishable Key für RPC-Aufrufe aus der App."],
        ],
        [1.85, 2.25, 2.4],
    )

    doc.add_heading("7. Aktueller Teststand", level=1)
    add_bullet(doc, "Supabase-Projekt SMART SnippetFlow ist verlinkt und die Migration wurde erfolgreich eingespielt.")
    add_bullet(doc, "Die Functions create-checkout-session und stripe-webhook sind deployed.")
    add_bullet(doc, "Ein Stripe-Testkauf hat die aktive Lizenz SSF-14F6-2E66-0817-032B erzeugt.")
    add_bullet(doc, "Die Supabase-RPC-Aktivierung wurde mit einem Test-Gerätehash erfolgreich geprüft und danach deaktiviert.")
    add_bullet(doc, "Die Desktop-App ist lokal mit Supabase URL und Publishable Key vorbereitet.")

    doc.add_heading("8. Betrieb und Support", level=1)
    add_matrix_table(
        doc,
        ["Frage", "Prüfung"],
        [
            ["Ist eine Lizenz bezahlt?", "Tabelle licenses: status active und passende Stripe Checkout Session ID."],
            ["Ist ein Gerät aktiviert?", "Tabelle license_activations: deactivated_at ist leer."],
            ["Warum kam keine Lizenz an?", "Tabelle stripe_events und license_audit_log prüfen; Stripe Webhook Deliveries ansehen."],
            ["Gerätelimit erreicht?", "Aktive Zeilen in license_activations zählen und alte Geräte deaktivieren."],
            ["Checkout funktioniert nicht?", "Stripe Secret Key und Price ID müssen aus demselben Stripe-Kontext stammen."],
        ],
        [2.2, 4.3],
    )

    doc.add_heading("9. Offene Ausbauschritte", level=1)
    add_bullet(doc, "Downloadbereich auf der Landingpage ergänzen.")
    add_bullet(doc, "Success-Seite nach der Zahlung bauen.")
    add_bullet(doc, "Lizenzschlüssel automatisch per E-Mail senden oder auf der Success-Seite anzeigen.")
    add_bullet(doc, "Produktionsbuild der Desktop-App mit Supabase URL und Publishable Key konfigurieren.")
    add_bullet(doc, "Optional: Stripe Customer Portal für Rechnungen, Zahlungsdaten und Kündigung ergänzen.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
